import sys
import os
import json
import win10toast
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QApplication, QComboBox, QDateEdit, QDialog, QDoubleSpinBox, QLineEdit
from tablegenerator_UI import Ui_TableGenerate
import generator
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer

date = datetime.now()
year = str(date.year)
month = str(date.month)
day = str(date.day)
now = year + month.rjust(2, '0') + day.rjust(2, '0')
program_file_path = os.path.dirname(os.path.abspath(__file__))

class MyMainForm(QDialog, Ui_TableGenerate):
    def __init__(self, parent=None) -> None:
        super(MyMainForm, self).__init__(parent=parent)
        self.setWindowTitle('Table Generator v1.2')
        self.setupUi(self)
        self.text_property = self.findChildren(QLineEdit)
        self.spin_property = self.findChildren(QDoubleSpinBox)
        self.date_property = self.findChildren(QDateEdit)
        self.combobox_property = self.findChildren(QComboBox)

        self.output_dir = os.getcwd()
        self.filename = now + '_' + '1'

        self.last_log_file = program_file_path + os.sep + 'last_info.json'
        if os.path.exists(self.last_log_file):
            self.load_info()
        if not self.Code_No.text():
            self.Code_No.setText('1'.rjust(9, '0'))
        
        self.OK_Button.clicked.connect(self.table_generate)
        self.toolButton.clicked.connect(self.open_path)
        self.calcButton.clicked.connect(self.calc_sum)
        self.pushButton_addmerge.clicked.connect(self.add_file_to_merge)
        self.pushButton_merge.clicked.connect(self.merge_many_to_one)

        self.files_to_merge = []

    @property
    def docx_filepath(self):
        return self.output_dir + '/' + self.filename + '.docx'

    @property
    def pdf_filepath(self):
        return self.output_dir + '/' + self.filename + '.pdf' 

    def save_info(self):

        dic = {}
        for obj in self.text_property:
            dic[obj.objectName()] = obj.text()
        for obj in self.spin_property:
            dic[obj.objectName()] = obj.value()
        for obj in self.date_property:
            dic[obj.objectName()] = obj.dateTime().toString('yyyy:MM:dd')
        for obj in self.combobox_property:
            dic[obj.objectName()] = obj.currentText()
        dic[self.dir_path.objectName()] = self.dir_path.text()

        dic['count'] = int(dic['count'])

        if dic['network_department'] == '其他...':
            dic['network_department'] = dic['network_department_other']
            dic.pop('network_department_other')
        
        if dic['start_add'] == '其他...':
            dic['start_add'] = dic['start_add_other']
            dic.pop('start_add_other')

        money = '%.2f' % dic['total_money']
        dic['total_money'] = str(money)
        num_list = generator.number_transfer(money)
        money_CN_list = ['money_penny','money_cent', 'money_one', 'money_ten', 'money_h', 'money_t', 'money_tt']
        for i in range(len(money_CN_list)):
            dic[money_CN_list[i]] = '' 
        i = -1
        for t in money_CN_list:
            try:
                dic[t] = num_list[i]
                i -= 1
            except IndexError:
                break

        with open('last_info.json', 'w') as f:
            f.write(json.dumps(dic))

        return dic

    def load_info(self):
        with open(self.last_log_file) as f:
            info_dict = json.load(f)

        info_dict['network_department_other'] = ''
        info_dict['start_add_other'] = ''
            
        for obj in self.text_property:
            try:
                obj.setText(info_dict[obj.objectName()])
            except KeyError:
                continue
        if not info_dict['Code_No']:
            info_dict['Code_No'] = '0'
        self.Code_No.setText(str(int(info_dict['Code_No']) + 1).rjust(9, '0'))
            
        for obj in self.spin_property:
            obj.setValue(float(info_dict[obj.objectName()]))
        for obj in self.combobox_property:
            obj.setCurrentText(info_dict[obj.objectName()])
        self.dir_path.setText(info_dict[self.dir_path.objectName()])

    def generate_filename(self):
        i = 1
        while True:
            self.filename = now + '_' + str(i)
            if os.path.exists(self.docx_filepath) or os.path.exists(self.pdf_filepath):
                i += 1
            else:
                break

    def open_path(self):
        file_dir= QtWidgets.QFileDialog.getExistingDirectory(self, "选择保存位置", os.getcwd())
        self.dir_path.setText(file_dir)
    
    def calc_sum(self):
        l = [self.package_value.value(),self.agency_fund_value.value(), self.self_fee_value.value(), self.transfer_fee_value.value(), self.delivery_cost_value.value()]
        sum_value = 0.0
        for i in l:
            i = float(i)
            sum_value += i
        self.total_money.setValue(sum_value)
    
    def table_generate(self):
        dic = self.save_info()
        if dic['dir_path']:
            self.output_dir = dic['dir_path']
        else:
            self.output_dir = os.getcwd()
        self.generate_filename()
        generator.table_generate(dic, self.docx_filepath)
        self._add_file_to_merge(self.docx_filepath)
        self.status_label.setText("表单文件 %s 已生成" % self.docx_filepath)
        current_No = int(self.Code_No.text())
        self.Code_No.setText(str(current_No + 1).rjust(9, '0'))
        
    def _add_file_to_merge(self, file):
        self.listWidget.addItem(file)
        self.files_to_merge.append(file)

    def add_file_to_merge(self):
        file_dir = QtWidgets.QFileDialog.getOpenFileName(self, '选择文件', os.getcwd())
        if file_dir[0]:
            self._add_file_to_merge(file_dir[0])
            self.status_label.setText("文件 %s 已添加" % file_dir[0].split('/')[-1])

    def merge_many_to_one(self):

        file_list = self.files_to_merge
        output_file = self.output_dir + '/' + now + '_总.docx'

        if len(file_list) == 0:
            return

        _page_break_docx = Document()
        _page_break_docx.add_page_break()
        # 保证样式一致
        composer = Composer(Document(file_list[0]))
        for file in file_list[1:]:
            #composer.append(_page_break_docx)
            composer.append(Document(file))
        composer.save(output_file)
        self.status_label.setText("表单文件 %s 已生成" % output_file)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    myWin = MyMainForm()
    myWin.show()
    sys.exit(app.exec_())
    
